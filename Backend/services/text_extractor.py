# services/text_extractor.py - Text Extraction Service
import os
from typing import Optional

# PDF text extraction using pypdf2
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word document text extraction  
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Text file extraction
import chardet

def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """Extract text from PDF file using PyPDF2"""
    if not PDF_AVAILABLE:
        print("PyPDF2 not available. PDF extraction disabled.")
        return None
    
    try:
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text.strip():
                        if page_num > 0:
                            text_content.append(f"\n--- Page {page_num + 1} ---\n")
                        text_content.append(page_text)
                        
                except Exception as page_error:
                    print(f"Failed to extract text from page {page_num + 1}: {page_error}")
                    continue
        
        if text_content:
            full_text = "\n".join(text_content)
            # Clean up extra whitespace
            lines = [line.strip() for line in full_text.split("\n") if line.strip()]
            return "\n".join(lines)
        
        return None
        
    except Exception as e:
        print(f"PDF text extraction failed for {file_path}: {e}")
        return None

def extract_text_from_docx(file_path: str) -> Optional[str]:
    """Extract text from Word document (.docx)"""
    if not DOCX_AVAILABLE:
        print("python-docx not available. DOCX extraction disabled.")
        return None
    
    try:
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        if text_content:
            return "\n".join(text_content)
        
        return None
        
    except Exception as e:
        print(f"DOCX text extraction failed for {file_path}: {e}")
        return None

def extract_text_from_txt(file_path: str) -> Optional[str]:
    """Extract text from plain text file with encoding detection"""
    try:
        # Detect file encoding
        with open(file_path, 'rb') as file:
            raw_data = file.read()
            encoding_result = chardet.detect(raw_data)
            encoding = encoding_result.get('encoding', 'utf-8')
        
        # Read file with detected encoding
        with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
            content = file.read()
            
            # Clean up the content
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            return '\n'.join(lines) if lines else None
            
    except Exception as e:
        print(f"TXT text extraction failed for {file_path}: {e}")
        return None

def extract_text_from_file(file_path: str, file_type: str) -> Optional[str]:
    """Main function to extract text from various file formats"""
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return None
    
    # Normalize file type
    file_type = file_type.lower().strip()
    
    try:
        if file_type == 'pdf':
            return extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return extract_text_from_docx(file_path)
        elif file_type == 'txt':
            return extract_text_from_txt(file_path)
        elif file_type == 'doc':
            print(f"Legacy DOC format not supported: {file_path}")
            return None
        else:
            print(f"Unsupported file type: {file_type}")
            return None
            
    except Exception as e:
        print(f"Text extraction failed for {file_path} ({file_type}): {e}")
        return None