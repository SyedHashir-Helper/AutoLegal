import json
import re
import os
import threading
import time
import uuid
from datetime import datetime, timedelta
from http.server import HTTPServer, SimpleHTTPRequestHandler
from urllib.parse import urlparse
import webbrowser
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn


class TemporaryFileServer:
    """Creates a temporary HTTP server to serve generated documents."""
    
    def __init__(self, port=8000):
        self.port = port
        self.server = None
        self.server_thread = None
        self.served_files = {}
        
    def add_file(self, file_path, expiry_minutes=30):
        """Add a file to be served temporarily."""
        file_id = str(uuid.uuid4())
        expiry_time = datetime.now() + timedelta(minutes=expiry_minutes)
        
        self.served_files[file_id] = {
            'path': file_path,
            'expiry': expiry_time,
            'original_name': os.path.basename(file_path)
        }
        
        return file_id
    
    def start_server(self):
        """Start the temporary HTTP server."""
        class CustomHandler(SimpleHTTPRequestHandler):
            def __init__(self, *args, **kwargs):
                self.server_instance = kwargs.pop('server_instance')
                super().__init__(*args, **kwargs)
            
            def do_GET(self):
                # Parse the URL to get file ID
                path = self.path.strip('/')
                
                if path.startswith('download/'):
                    file_id = path.split('/')[-1]
                    self.handle_download(file_id)
                elif path == '' or path == 'index.html':
                    self.send_file_list()
                else:
                    self.send_error(404, "File not found")
            
            def handle_download(self, file_id):
                """Handle file download requests."""
                if file_id not in self.server_instance.served_files:
                    self.send_error(404, "File not found or expired")
                    return
                
                file_info = self.server_instance.served_files[file_id]
                
                # Check if file has expired
                if datetime.now() > file_info['expiry']:
                    del self.server_instance.served_files[file_id]
                    self.send_error(410, "File has expired")
                    return
                
                # Check if file exists
                if not os.path.exists(file_info['path']):
                    self.send_error(404, "File not found on disk")
                    return
                
                # Serve the file
                try:
                    with open(file_info['path'], 'rb') as f:
                        content = f.read()
                    
                    self.send_response(200)
                    self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    self.send_header('Content-Disposition', f'attachment; filename="{file_info["original_name"]}"')
                    self.send_header('Content-Length', str(len(content)))
                    self.end_headers()
                    self.wfile.write(content)
                    
                except Exception as e:
                    self.send_error(500, f"Error serving file: {str(e)}")
            
            def send_file_list(self):
                """Send a simple HTML page with available downloads."""
                html_content = self.generate_download_page()
                
                self.send_response(200)
                self.send_header('Content-Type', 'text/html; charset=utf-8')
                self.send_header('Content-Length', str(len(html_content.encode('utf-8'))))
                self.end_headers()
                self.wfile.write(html_content.encode('utf-8'))
            
            def generate_download_page(self):
                """Generate HTML page with download links."""
                current_time = datetime.now()
                active_files = {k: v for k, v in self.server_instance.served_files.items() 
                              if v['expiry'] > current_time}
                
                html = """
                <!DOCTYPE html>
                <html>
                <head>
                    <title>Document Downloads</title>
                    <style>
                        body { font-family: Arial, sans-serif; margin: 40px; }
                        .file-item { margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 5px; }
                        .download-btn { background: #007cba; color: white; padding: 10px 20px; text-decoration: none; border-radius: 3px; }
                        .download-btn:hover { background: #005a87; }
                        .expiry { color: #666; font-size: 0.9em; }
                    </style>
                </head>
                <body>
                    <h1>Generated Documents</h1>
                """
                
                if active_files:
                    for file_id, file_info in active_files.items():
                        expiry_str = file_info['expiry'].strftime('%Y-%m-%d %H:%M:%S')
                        html += f"""
                        <div class="file-item">
                            <h3>{file_info['original_name']}</h3>
                            <p class="expiry">Expires: {expiry_str}</p>
                            <a href="/download/{file_id}" class="download-btn">Download</a>
                        </div>
                        """
                else:
                    html += "<p>No documents available for download.</p>"
                
                html += """
                    <hr>
                    <p><small>This server will automatically shut down when all files expire.</small></p>
                </body>
                </html>
                """
                return html
            
            def log_message(self, format, *args):
                """Override to reduce server logging."""
                pass
        
        def handler_factory(*args, **kwargs):
            return CustomHandler(*args, server_instance=self, **kwargs)
        
        try:
            self.server = HTTPServer(('localhost', self.port), handler_factory)
            self.server_thread = threading.Thread(target=self.server.serve_forever, daemon=True)
            self.server_thread.start()
            return True
        except Exception as e:
            print(f"Failed to start server on port {self.port}: {e}")
            return False
    
    def stop_server(self):
        """Stop the HTTP server."""
        if self.server:
            self.server.shutdown()
            self.server.socket.close()
    
    def cleanup_expired_files(self):
        """Remove expired files from the served files list."""
        current_time = datetime.now()
        expired_files = [k for k, v in self.served_files.items() if v['expiry'] <= current_time]
        
        for file_id in expired_files:
            file_path = self.served_files[file_id]['path']
            del self.served_files[file_id]
            
            # Optionally delete the actual file
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception as e:
                print(f"Warning: Could not delete expired file {file_path}: {e}")
    
    def get_download_url(self, file_id):
        """Get the download URL for a file."""
        return f"http://localhost:{self.port}/download/{file_id}"
    
    def get_server_url(self):
        """Get the main server URL."""
        return f"http://localhost:{self.port}"


class DocumentGenerator:
    def __init__(self, combined_json_file=None, combined_json_data=None, enable_server=True):
        """Initialize the document generator with combined JSON containing structure and input data."""
        if combined_json_data:
            self.combined_data = combined_json_data
        elif combined_json_file:
            self.combined_data = self.load_json(combined_json_file)
        else:
            raise ValueError("Either combined_json_file or combined_json_data must be provided")
        
        # Extract structure and input data from combined JSON
        self.structure = self.combined_data.get('structure', {})
        self.input_data = self.combined_data.get('input_data', {})
        self.document_type = self.combined_data.get('document_type', 'Document')
        
        self.doc = Document()
        self.enable_server = enable_server
        self.file_server = None
        
        if self.enable_server:
            self.file_server = TemporaryFileServer()
            if not self.file_server.start_server():
                print("Warning: Could not start file server. Download URLs will not be available.")
                self.file_server = None
        
        self.setup_document()
    
    def load_json(self, file_path):
        """Load JSON data from file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def setup_document(self):
        """Set up document formatting based on structure JSON."""
        formatting = self.structure.get('formatting', {})
        
        # Set margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def replace_placeholders(self, text):
        """Replace placeholders in text with actual values from input data."""
        if not isinstance(text, str):
            return text
        
        # Find all placeholders in the format {{key}} or {{nested.key}}
        placeholders = re.findall(r'\{\{([^}]+)\}\}', text)
        
        for placeholder in placeholders:
            value = self.get_nested_value(self.input_data, placeholder)
            if value is not None:
                text = text.replace(f'{{{{{placeholder}}}}}', str(value))
        
        return text
    
    def get_nested_value(self, data, key_path):
        """Get value from nested dictionary using dot notation."""
        keys = key_path.split('.')
        current = data
        
        for key in keys:
            if isinstance(current, dict) and key in current:
                current = current[key]
            else:
                return None
        
        return current
    
    def add_title(self, title_config):
        """Add document title."""
        title = self.doc.add_heading(level=1)
        title_run = title.runs[0] if title.runs else title.add_run()
        title_run.text = self.replace_placeholders(title_config['content'])
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_header_section(self, header_config):
        """Add header section with document details (used in SOW)."""
        if header_config.get('type') == 'header_details':
            for item in header_config.get('content', []):
                header_para = self.doc.add_paragraph()
                label = self.replace_placeholders(item.get('label', ''))
                value = self.replace_placeholders(item.get('value', ''))
                
                if label and value:
                    label_run = header_para.add_run(f"{label} ")
                    label_run.bold = True
                    header_para.add_run(value)
    
    def add_parties_section(self, parties_config):
        """Add parties section with party details (flexible for different party types)."""
        # Add intro paragraph
        intro_para = self.doc.add_paragraph()
        intro_para.add_run(self.replace_placeholders(parties_config['content']))
        
        # Handle flexible party structure
        for party_key, party_config in parties_config.get('subsections', {}).items():
            if party_config.get('type') == 'party_details':
                party_para = self.doc.add_paragraph()
                party_para.add_run(self.replace_placeholders(party_config['label'])).bold = True
                party_para.add_run('\n')
                
                for field_key, field_value in party_config.get('fields', {}).items():
                    field_content = self.replace_placeholders(field_value)
                    if field_content and field_content.strip():  # Only add non-empty fields
                        party_para.add_run(field_content + '\n')
    
    def add_section(self, section_config):
        """Add a regular section with title and content."""
        # Add section heading
        heading = self.doc.add_heading(level=2)
        section_title = f"{section_config['number']}. {section_config['title']}"
        heading.add_run(section_title)
        
        # Add section content
        content_para = self.doc.add_paragraph()
        content_para.add_run(self.replace_placeholders(section_config['content']))
        
        # Add subsections if they exist
        if 'subsections' in section_config:
            self.add_subsections(section_config['subsections'])
    
    def add_subsections(self, subsections_config):
        """Add subsections, including bullet lists, timeline details, and payment details."""
        subsection_type = subsections_config.get('type')
        
        if subsection_type == 'bullet_list':
            # Handle bullet list items
            for item in subsections_config.get('items', []):
                bullet_para = self.doc.add_paragraph()
                bullet_para.style = 'List Bullet'
                bullet_para.add_run(self.replace_placeholders(item.get('content', '')))
                
        elif subsection_type == 'timeline_details' or subsection_type == 'payment_details':
            # Handle label-value pairs (like timeline and payment details)
            for item in subsections_config.get('items', []):
                detail_para = self.doc.add_paragraph()
                label = self.replace_placeholders(item.get('label', ''))
                value = self.replace_placeholders(item.get('value', ''))
                
                # Format as "Label: Value"
                if label and value:
                    label_run = detail_para.add_run(f"{label} ")
                    label_run.bold = True
                    detail_para.add_run(value)
                elif label:
                    detail_para.add_run(label)
                elif value:
                    detail_para.add_run(value)
                    
        elif subsection_type == 'header_details':
            # Handle header details (for SOW header section)
            for item in subsections_config:
                header_para = self.doc.add_paragraph()
                label = self.replace_placeholders(item.get('label', ''))
                value = self.replace_placeholders(item.get('value', ''))
                
                if label and value:
                    label_run = header_para.add_run(f"{label} ")
                    label_run.bold = True
                    header_para.add_run(value)
        
        else:
            # Fallback for unknown subsection types
            items = subsections_config.get('items', [])
            for item in items:
                para = self.doc.add_paragraph()
                if isinstance(item, dict):
                    if 'content' in item:
                        para.add_run(self.replace_placeholders(item['content']))
                    elif 'label' in item and 'value' in item:
                        label = self.replace_placeholders(item['label'])
                        value = self.replace_placeholders(item['value'])
                        label_run = para.add_run(f"{label} ")
                        label_run.bold = True
                        para.add_run(value)
                else:
                    para.add_run(self.replace_placeholders(str(item)))
    
    def add_signature_section(self, signature_config):
        """Add signature section with flexible party signatures."""
        # Add witness clause if present
        if signature_config.get('content'):
            witness_para = self.doc.add_paragraph()
            witness_para.add_run(signature_config['content']).bold = True
            # Add spacing
            self.doc.add_paragraph()
        
        # Add signature blocks for all parties
        for party_key, party_sig in signature_config.get('signatures', {}).items():
            sig_para = self.doc.add_paragraph()
            
            # Add party label (bold)
            party_label = self.replace_placeholders(party_sig['label'])
            sig_para.add_run(party_label).bold = True
            sig_para.add_run('\n')
            
            # Add signature line
            sig_para.add_run('_' * 50 + '\n')
            
            # Add signature fields without any automatic prefixes
            for field_key, field_value in party_sig.get('fields', {}).items():
                field_content = self.replace_placeholders(field_value)
                
                if field_content and field_content.strip():  # Only add non-empty fields
                    # Add content exactly as provided, no additional formatting
                    sig_para.add_run(field_content + '\n')
            
            # Add spacing between signature blocks
            self.doc.add_paragraph()
    
    def generate_document(self, output_file=None, expiry_minutes=30):
        """Generate the complete document and create download URL."""
        structure_data = self.structure
        
        # Generate unique filename if not provided
        if output_file is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            doc_type = self.document_type.replace(' ', '_')
            output_file = f"{doc_type}_{timestamp}.docx"
        
        # Add title
        if 'title' in structure_data:
            self.add_title(structure_data['title'])
        
        # Add header section (for SOW documents)
        if 'header_section' in structure_data:
            self.add_header_section(structure_data['header_section'])
            # Add spacing after header
            self.doc.add_paragraph()
        
        # Add parties section
        if 'parties' in structure_data:
            self.add_parties_section(structure_data['parties'])
        
        # Add all sections
        if 'sections' in structure_data:
            for section in structure_data['sections']:
                self.add_section(section)
        
        # Add signature section
        if 'signature_section' in structure_data:
            self.add_signature_section(structure_data['signature_section'])
        
        # Save document
        self.doc.save(output_file)
        print(f"✅ Document generated successfully: {output_file}")
        
        # Create download URL if server is enabled
        download_info = None
        if self.file_server:
            try:
                file_id = self.file_server.add_file(output_file, expiry_minutes)
                download_url = self.file_server.get_download_url(file_id)
                server_url = self.file_server.get_server_url()
                
                download_info = {
                    'file_path': output_file,
                    'download_url': download_url,
                    'server_url': server_url,
                    'expiry_minutes': expiry_minutes,
                    'file_id': file_id
                }
                
                print(f"🌐 Download URL: {download_url}")
                print(f"🌐 Server Dashboard: {server_url}")
                print(f"⏰ Link expires in {expiry_minutes} minutes")
                
                # Optionally open browser
                try:
                    webbrowser.open(server_url)
                    print("🚀 Opening browser...")
                except Exception as e:
                    print(f"Could not open browser automatically: {e}")
                    
            except Exception as e:
                print(f"Warning: Could not create download URL: {e}")
        
        return download_info
    
    def cleanup_and_stop_server(self):
        """Clean up expired files and stop the server."""
        if self.file_server:
            self.file_server.cleanup_expired_files()
            if not self.file_server.served_files:  # No more files to serve
                self.file_server.stop_server()
                print("🛑 Server stopped - no more files to serve")


def main():
    """Main function to run the document generator."""
    try:
        # Initialize generator with server enabled
        print("🚀 Starting Document Generator...")
        generator = DocumentGenerator(combined_json_file='combined_document.json', enable_server=True)
        
        # Generate document with download URL
        download_info = generator.generate_document(expiry_minutes=60)  # 1 hour expiry
        
        if download_info:
            print("\n" + "="*50)
            print("📄 DOCUMENT GENERATED SUCCESSFULLY")
            print("="*50)
            print(f"📁 File: {download_info['file_path']}")
            print(f"🔗 Download: {download_info['download_url']}")
            print(f"🌐 Dashboard: {download_info['server_url']}")
            print(f"⏰ Expires: {download_info['expiry_minutes']} minutes")
            print("="*50)
            
            # Keep the server running
            print("\n💡 Server is running. Press Ctrl+C to stop.")
            try:
                while True:
                    time.sleep(10)
                    generator.file_server.cleanup_expired_files()
                    
                    # Stop server if no files left
                    if not generator.file_server.served_files:
                        print("⏰ All files have expired. Stopping server.")
                        break
                        
            except KeyboardInterrupt:
                print("\n🛑 Stopping server...")
            finally:
                generator.cleanup_and_stop_server()
        
    except FileNotFoundError as e:
        print(f"❌ Error: {e}")
        print("Make sure 'combined_document.json' file exists in the current directory.")
    except json.JSONDecodeError as e:
        print(f"❌ Error parsing JSON file: {e}")
    except Exception as e:
        print(f"❌ An error occurred: {e}")


def generate_document_only(combined_json_file='combined_document.json', combined_json_data=None, output_file=None):
    """Generate document without starting server (for programmatic use)."""
    try:
        generator = DocumentGenerator(
            combined_json_file=combined_json_file, 
            combined_json_data=combined_json_data, 
            enable_server=False
        )
        
        if output_file is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            doc_type = generator.document_type.replace(' ', '_')
            output_file = f"{doc_type}_{timestamp}.docx"
        
        generator.generate_document(output_file)
        return output_file
        
    except Exception as e:
        print(f"❌ Error generating document: {e}")
        return None


def start_server_with_existing_file(file_path, expiry_minutes=30, port=8000):
    """Start server to serve an existing document file."""
    try:
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return None
        
        server = TemporaryFileServer(port=port)
        if server.start_server():
            file_id = server.add_file(file_path, expiry_minutes)
            download_url = server.get_download_url(file_id)
            server_url = server.get_server_url()
            
            print(f"🌐 Download URL: {download_url}")
            print(f"🌐 Server Dashboard: {server_url}")
            print(f"⏰ Link expires in {expiry_minutes} minutes")
            
            # Open browser
            try:
                webbrowser.open(server_url)
            except:
                pass
            
            return {
                'server': server,
                'download_url': download_url,
                'server_url': server_url,
                'file_id': file_id
            }
        else:
            print(f"❌ Failed to start server on port {port}")
            return None
            
    except Exception as e:
        print(f"❌ Error starting server: {e}")
        return None


if __name__ == "__main__":
    main()